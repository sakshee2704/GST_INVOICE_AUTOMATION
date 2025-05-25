from docx import Document
from datetime import date,datetime,timedelta
import pandas as pd
from docx2pdf import convert
import numpy as np
import pikepdf
from email_with_attachment import send_email_with_attachment


#Function to replace placeholders in the template
# This function recursively searches and replaces placeholders in paragraphs and tables
def replace_placeholders(doc, placeholders):
   for paragraph in doc.paragraphs:
      for placeholder, value in placeholders.items():
           if placeholder in paragraph.text:
               for run in paragraph.runs:
                  run.text = run.text.replace(placeholder, str(value)) # Replace placeholder with actual value
      for table in doc.tables: # Process tables within the document
        for row in table.rows:
           for cell in row.cells:
               replace_placeholders(cell, placeholders)

# Function to handle paths
def get_paths():
    template_path = "C:/Users/Admin/OneDrive/Desktop/HDFL_Invoice/HDFC_Life/HDFC/Endorsement_II_July_2024.docx"
    output_path = "C:/Users/Admin/OneDrive/Desktop/HDFL_Invoice/HDFC_Life/" #HDFC
    return template_path, output_path

# Function to convert Excel serial date format to standard date
def excel_serial_to_date(serial_number):
    base_date = datetime(1899, 12, 30)  # Excel's base date
    # Convert numpy int64 to standard Python int if needed
    if isinstance(serial_number, np.int64):
        serial_number = int(serial_number)
    actual_date = base_date + timedelta(days=serial_number)
    return actual_date.strftime("%m%y")
                                    
# Define the main function
def generate_invoice(invoice_details):
    template_path, output_path = get_paths()

    todays_date= str(date.today().strftime("%d/%m/%Y")) # Get today's date
    # Load the template document
    doc = Document(template_path)
    template_date=invoice_details['ISSDATE']
    template_date = np.int64(template_date)
    time_date=excel_serial_to_date(template_date)
    invoice_details['ISSDATE']=time_date

    # Concatenate GST invoice number with additional details
    Gst_invoice_no=invoice_details['GST_INVOICE_NO']
    sequenceNo=invoice_details['SequenceNo']
    concat_gst=f"{Gst_invoice_no}/{time_date}/{sequenceNo:04d}"
    invoice_details['GST_INVOICE_NO']=concat_gst
    invoice_details['LIFE_INSURANCE']="Life Insurance Services"

    # Define the data to replace the placeholders
    placeholders = {
        "<CONTRACT_NO>": invoice_details['CONTRACT_NO'],
        "<OWNER_CLIENT_NO>": invoice_details['OWNER_CLIENT_NO'],                       
        "<CLIENT_FULLNAME>": invoice_details['CLIENT_FULLNAME'],
        "<INSURED_FROM>": invoice_details['INSURED_FROM(RCD date)'],
        "<DATE>": todays_date,
        "<GSTIN no>": invoice_details['GSTIN no'],
        "InvoiceNo": invoice_details['GST_INVOICE_NO'],
        "<ZGSTNO>": invoice_details['ZGSTNO'],
        "<CLTADDR01>": invoice_details['CLTADDR01'],
        "<CLTADDR02>": invoice_details['CLTADDR02'],
        "<CLTADDR03>": invoice_details['CLTADDR03'],
        "<CLTADDR04>": invoice_details['CLTADDR04'],
        "<CLTADDR05>":  invoice_details['CLTADDR05'],
        "<CLTPCODE>":  invoice_details['CLTADDR05'],
        "<CLTPCODE>":  str(invoice_details['ISSDATE']),
        "DD/MM/YYYY": str(invoice_details['ISSDATE']),
        "DD/MM/YYYY": str(invoice_details['ISSDATE']),
        "DD/MM/YYYY": str(invoice_details['ISSDATE']),
        "PARTICULARS_1": str(invoice_details['LIFE_INSURANCE']),
        "PREMIUM_1": invoice_details['NETPREMIUM'],
    }

    # Calculate totals for GST and other amounts
    total=0
    if invoice_details.get('ZHIGSTAMT'):   #Check for IGST(Column L)
        placeholders["gst"] = "IGST"
        placeholders["XXX"] = invoice_details['ZHIGSTAMT']
        total=invoice_details['ZHIGSTAMT']
    elif invoice_details.get('ZHSGSTAMT') and not invoice_details.get('ZHIGSTAMT'):   #Check for CGST(In Column J)
        placeholders["gst"] = "CGST"
        placeholders["XXX"] = invoice_details['ZHSGSTAMT']
        total=invoice_details['ZHSGSTAMT']
    
    placeholders["XX"] = invoice_details['ZHCGSTAMT']
    placeholders["<Total>"] = round(total+invoice_details['ZHCGSTAMT']+invoice_details['NETPREMIUM'],2)
    
     # Generate a filename using the GST number
    GST_no_generated=invoice_details['GST_INVOICE_NO']
    filename=GST_no_generated.replace('/', '_')
    
    # Replace the placeholders with actual data
    replace_placeholders(doc, placeholders)
    # Save the modified document with a new name output_path+=filename+".docx"
    doc.save(output_path+filename+".docx")
    print("Data has been inserted into the template!")

    # Convert to PDF
    convert(output_path+filename+".docx", output_path+filename+".pdf")
    
    # Encrypt the PDF with a password
    def encrypt_pdf(input_pdf_path, output_pdf_path, password):
    # Open the original PDF
     with pikepdf.open(input_pdf_path) as pdf:
        # Encrypt the PDF with the given password
        pdf.save(output_pdf_path, encryption=pikepdf.Encryption(owner=password, user=password, R=4))
    
    date_password = invoice_details['INSURED_FROM(RCD date)']
    
    # Ensure date_password is a datetime object
    if isinstance(date_password, str):  
     date_password = datetime.strptime(date_password, '%d/%m/%Y')  # Adjusted format to match input

    password = date_password.strftime('%d%m')  # Returns a string in "DDMM" format
    # print(password)

    Pass_files="C:/Users/Admin/OneDrive/Desktop/HDFL_Invoice/HDFC_Life/HDFC/final_files/"
    encrypt_pdf(output_path+filename+".pdf", Pass_files+filename+".pdf", password)

    # Send the encrypted file via email
    send_email_with_attachment(Pass_files+filename+".pdf")

